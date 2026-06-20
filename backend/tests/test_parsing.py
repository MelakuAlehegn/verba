import io

import docx
import pytest

from app.services.rag.parsing import DocumentParseError, extract_text

PARSING_MODULE = "app.services.rag.parsing"


def test_extract_text_file_decodes_markdown() -> None:
    text = extract_text(b"# Title\n\nA paragraph.", filename="notes.md")
    assert "Title" in text
    assert "A paragraph." in text


def test_extract_docx_round_trip() -> None:
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text(buffer.getvalue(), filename="report.docx")

    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_extract_pdf_joins_page_text(monkeypatch) -> None:
    class FakePage:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class FakeReader:
        def __init__(self, *_args, **_kwargs) -> None:
            self.pages = [FakePage("Page one."), FakePage("Page two.")]

    monkeypatch.setattr(f"{PARSING_MODULE}.PdfReader", FakeReader)

    text = extract_text(b"%PDF-fake", filename="doc.pdf")

    assert "Page one." in text
    assert "Page two." in text


def test_unsupported_type_raises() -> None:
    with pytest.raises(DocumentParseError):
        extract_text(b"data", filename="sheet.csv")


def test_empty_text_raises() -> None:
    with pytest.raises(DocumentParseError):
        extract_text(b"   \n  ", filename="blank.txt")
